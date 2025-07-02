import os
import shutil
import hashlib
from pathlib import Path
from datetime import datetime
import customtkinter as ctk
from tkinter import filedialog
import webbrowser
import threading

from PyPDF2 import PdfReader
from docx import Document

ctk.set_appearance_mode("System")
ctk.set_default_color_theme("blue")


def file_hash(filepath):
    hasher = hashlib.sha256()
    with open(filepath, "rb") as f:
        while chunk := f.read(8192):
            hasher.update(chunk)
    return hasher.hexdigest()


def get_doc_date(path):
    try:
        if path.suffix.lower() == ".pdf":
            reader = PdfReader(str(path))
            raw_date = reader.metadata.get("/CreationDate", "")
            if raw_date.startswith("D:"):
                return datetime.strptime(raw_date[2:10], "%Y%m%d").strftime("%Y-%m-%d")
        elif path.suffix.lower() == ".docx":
            doc = Document(str(path))
            props = doc.core_properties
            if props.created:
                return props.created.strftime("%Y-%m-%d")
    except Exception:
        pass
    return None


def get_unique_name(base_path, filename, suffix=""):
    name, ext = os.path.splitext(filename)
    candidate = Path(base_path) / f"{name}{suffix}{ext}"
    i = 1
    while candidate.exists():
        candidate = Path(base_path) / f"{name}{suffix}_{i}{ext}"
        i += 1
    return candidate


def get_text_snippet(path):
    try:
        if path.suffix.lower() == ".txt":
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.readline().strip()
        elif path.suffix.lower() == ".docx":
            doc = Document(str(path))
            return doc.paragraphs[0].text.strip() if doc.paragraphs else ""
        elif path.suffix.lower() == ".pdf":
            reader = PdfReader(str(path))
            if reader.pages:
                return reader.pages[0].extract_text().strip().split("\n")[0][:200]
    except Exception:
        pass
    return ""


def collect_documents(source_dir, target_dir, update_status, update_progress, dry_run=False):
    exts = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".txt", ".csv", ".odt", ".rtf"}
    hashes = {}
    copied = 0
    renamed = 0
    log = []

    all_files = [
        Path(root) / f for root, _, files in os.walk(source_dir) for f in files
        if os.path.splitext(f)[1].lower() in exts
    ]
    total = len(all_files)

    for i, src_path in enumerate(all_files):
        update_progress(int((i + 1) / total * 100))
        update_status(f"Processing {src_path.name} ({i+1}/{total})")

        file_h = file_hash(src_path)
        date_folder = get_doc_date(src_path) or datetime.fromtimestamp(src_path.stat().st_mtime).strftime("%Y-%m-%d")
        if not date_folder:
            date_folder = "brak_dat"
        target_subdir = target_dir / date_folder

        if not dry_run:
            os.makedirs(target_subdir, exist_ok=True)

        snippet = get_text_snippet(src_path)

        if file_h in hashes:
            new_dst = get_unique_name(target_subdir, src_path.name, suffix="_dup")
            log.append(f"DUPLICATE: {src_path} -> {new_dst.relative_to(target_dir)} | {snippet}")
            if not dry_run:
                shutil.copy2(src_path, new_dst)
            renamed += 1
            continue

        dst_path = get_unique_name(target_subdir, src_path.name)
        if dst_path.name != src_path.name:
            log.append(f"RENAME: {src_path.name} -> {dst_path.relative_to(target_dir)} | {snippet}")
            renamed += 1
        else:
            log.append(f"COPY: {src_path.name} -> {dst_path.relative_to(target_dir)} | {snippet}")

        if not dry_run:
            shutil.copy2(src_path, dst_path)
        hashes[file_h] = dst_path.name
        copied += 1

    if not dry_run:
        os.makedirs(target_dir, exist_ok=True)
        log_path = target_dir / "log.txt"
        with open(log_path, "w", encoding="utf-8") as f:
            f.write(f"Log generated at: {datetime.now()}\n")
            f.write(f"Source folder: {source_dir}\nDestination folder: {target_dir}\nDry run: {dry_run}\n\n")
            for line in log:
                f.write(line + "\n")
            f.write(f"\nFiles processed: {copied}\nDuplicates renamed: {renamed}\n")

    return copied, renamed, target_dir


class DocumentCollectorApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Document Collector")
        self.geometry("500x400")

        self.label = ctk.CTkLabel(self, text="This tool scans a folder (recursively),\ncollects documents and saves them to a new folder on your Desktop.\nDuplicates are renamed, not skipped.", justify="center")
        self.label.pack(pady=20)

        self.dry_run_var = ctk.BooleanVar()
        self.dry_run_check = ctk.CTkCheckBox(self, text="Dry run (simulate only)", variable=self.dry_run_var)
        self.dry_run_check.pack(pady=5)

        self.select_button = ctk.CTkButton(self, text="Select Folder", command=self.select_folder)
        self.select_button.pack(pady=10)

        self.status_label = ctk.CTkLabel(self, text="")
        self.status_label.pack(pady=10)

        self.progress = ctk.CTkProgressBar(self)
        self.progress.set(0)
        self.progress.pack(pady=10, fill="x", padx=20)

    def select_folder(self):
        folder = filedialog.askdirectory(title="Select folder to scan")
        if not folder:
            return

        self.status_label.configure(text="Processing... Please wait.")
        self.update_idletasks()

        now = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        desktop = Path.home() / "Desktop"
        target_dir = desktop / f"COLLECTED_DOCUMENTS_{now}"

        dry_run = self.dry_run_var.get()

        def run():
            copied, renamed, target = collect_documents(
                folder, target_dir,
                update_status=self.status_label.configure,
                update_progress=lambda v: self.progress.set(v / 100),
                dry_run=dry_run
            )
            self.status_label.configure(text="Done!")
            self.progress.set(1)
            self.after(100, lambda: SummaryWindow(self, copied, renamed, target, dry_run))

        threading.Thread(target=run).start()


class SummaryWindow(ctk.CTkToplevel):
    def __init__(self, parent, copied, renamed, target_dir, dry_run):
        super().__init__(parent)
        self.title("Summary")
        self.geometry("500x280")
        self.grab_set()

        if dry_run:
            preface = "🧪 DRY RUN MODE: No files were actually copied. This was a simulation.\n"
        else:
            preface = "✅ OPERATION COMPLETE\n"

        msg = (
            f"{preface}\n"
            f"Files scanned: {copied + renamed}\n"
            f"Unique files: {copied}\n"
            f"Duplicates renamed: {renamed}\n"
            f"Destination folder: {target_dir}\n"
        )

        ctk.CTkLabel(self, text=msg, justify="left", wraplength=450).pack(pady=20)

        if not dry_run:
            ctk.CTkButton(self, text="Open Folder", command=lambda: webbrowser.open(str(target_dir))).pack(pady=5)

        ctk.CTkButton(self, text="Close", command=self.destroy).pack(pady=5)


if __name__ == "__main__":
    app = DocumentCollectorApp()
    app.mainloop()
